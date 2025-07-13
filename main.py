import streamlit as st #python frameworkturns your Python scripts into interactive web apps
import fitz  # PyMuPDF,to open, manipulate, and extract information from PDF files
import pandas as pd
from docx import Document
from io import BytesIO
import google.generativeai as genai #to interact with google's llm's
from dotenv import load_dotenv
import os

# Load Gemini API key
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))  #to set the api key

# PDF extraction
def extract_pdf_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    pages = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        pages.append({"Page": page_num, "Content": text.strip()})
    return pages

# Gemini summary
def summarize_text(text):
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(f"Summarize this:\n\n{text}")
        return response.text
    except Exception as e:
        return f"Error: {e}"
# Word file generator
def generate_word(data, summarize):
    doc = Document()  #empty word doc.
    doc.add_heading("PDF to Word Export", 0)
    for entry in data:
        doc.add_heading(f"Page {entry['Page']}", level=1)
        doc.add_paragraph(entry["Content"])
        if summarize and "Summary" in entry:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(entry["Summary"])
    buffer = BytesIO()   #save the doc. in memory
    doc.save(buffer)
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "output.docx"

# Excel generator
def generate_excel(data):
    df = pd.DataFrame(data)
    buffer = BytesIO()
    df.to_excel(buffer, index=False, engine='openpyxl')
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "output.xlsx"

# CSV generator
def generate_csv(data):
    df = pd.DataFrame(data)
    buffer = BytesIO()
    buffer.write(df.to_csv(index=False).encode("utf-8"))
    buffer.seek(0)
    return buffer, "text/csv", "output.csv"


# GUI with Streamlit
st.title("📄 PDF to WORD/CSV/XLSX Converter (with optional Gemini Summary)")

uploaded_pdf = st.file_uploader("Upload a PDF file", type="pdf")

summarize = st.checkbox("Summarize each page using Gemini?", value=False)
output_format = st.radio("Select output format", ["Word (.docx)", "Excel (.xlsx)", "CSV (.csv)"])

if uploaded_pdf and st.button("Convert"):
    with st.spinner("Extracting text..."):
        data = extract_pdf_text(uploaded_pdf)
        if summarize:
            st.info("Summarizing pages with Gemini...")
            for entry in data:
                entry["Summary"] = summarize_text(entry["Content"])

        if output_format == "Word (.docx)":
            buffer, mime, filename = generate_word(data, summarize)
        elif output_format == "Excel (.xlsx)":
            buffer, mime, filename = generate_excel(data)
        else:
            buffer, mime, filename = generate_csv(data)

        # Download button
        st.success("✅ File ready!")
        st.download_button("⬇ Download File", buffer, file_name=filename, mime=mime)